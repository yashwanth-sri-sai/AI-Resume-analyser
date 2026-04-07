"""
utils/file_parser.py
Handles extraction of text from PDF and DOCX resume files.
"""

import io
import re
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
        return "\n".join(text_pages)
    except ImportError:
        return _fallback_pdf_extract(file_bytes)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _fallback_pdf_extract(file_bytes: bytes) -> str:
    """Minimal PDF text extraction without pdfplumber."""
    try:
        text = file_bytes.decode("latin-1", errors="ignore")
        # Extract text between stream markers (very basic)
        streams = re.findall(rb"stream\r?\n(.*?)\r?\nendstream", file_bytes, re.DOTALL)
        readable = []
        for s in streams:
            try:
                decoded = s.decode("latin-1", errors="ignore")
                # Keep only printable ASCII
                clean = re.sub(r"[^\x20-\x7E\n]", " ", decoded)
                if len(clean.strip()) > 20:
                    readable.append(clean.strip())
            except Exception:
                pass
        return "\n".join(readable) if readable else "Could not extract text from PDF."
    except Exception:
        return "Could not extract text from PDF."


def parse_resume_file(uploaded_file) -> dict:
    """
    Main entry point for parsing a Streamlit UploadedFile.
    Returns a dict with 'text', 'file_type', 'file_name', 'file_size'.
    """
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name
    file_size = len(file_bytes)
    file_type = "unknown"

    if file_name.lower().endswith(".pdf"):
        file_type = "pdf"
        text = extract_text_from_pdf(file_bytes)
    elif file_name.lower().endswith(".docx"):
        file_type = "docx"
        text = extract_text_from_docx(file_bytes)
    elif file_name.lower().endswith(".doc"):
        file_type = "doc"
        raise ValueError(".doc files are not supported. Please convert to .docx or .pdf")
    elif file_name.lower().endswith(".txt"):
        file_type = "txt"
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {file_name}")

    if not text or len(text.strip()) < 50:
        raise ValueError("Could not extract sufficient text from the file. The file may be image-based or corrupted.")

    return {
        "text": text,
        "file_type": file_type,
        "file_name": file_name,
        "file_size": file_size,
        "char_count": len(text),
        "word_count": len(text.split()),
    }


def validate_file(uploaded_file) -> tuple[bool, str]:
    """Validate file type and size before processing."""
    if uploaded_file is None:
        return False, "No file uploaded."

    name = uploaded_file.name.lower()
    if not any(name.endswith(ext) for ext in [".pdf", ".docx", ".txt"]):
        return False, "Only PDF, DOCX, and TXT files are supported."

    # Check file size (max 10MB)
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset file pointer
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > 10:
        return False, f"File too large ({size_mb:.1f}MB). Maximum size is 10MB."

    return True, "File is valid."
